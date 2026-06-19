import os
import shutil
from datetime import datetime
from docx import Document
from src.database import DatabaseManager
from src.models import Entrevista, Sujeito

class ImportService:
    @staticmethod
    def extract_text_from_txt(filepath):
        encodings = ['utf-8', 'latin-1', 'cp1252', 'utf-16']
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise ValueError("Não foi possível decodificar o arquivo TXT usando codificações comuns (UTF-8, Latin-1, etc.).")

    @staticmethod
    def extract_text_from_docx(filepath):
        try:
            doc = Document(filepath)
            paragraphs = [p.text for p in doc.paragraphs]
            
            # Também extrai texto de tabelas se houver
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs.append(cell.text)
                        
            return "\n".join(paragraphs)
        except Exception as e:
            raise ValueError(f"Erro ao ler arquivo DOCX: {str(e)}")

    @classmethod
    def import_interview(cls, filepath, project_root_path, sujeito_id=None):
        """
        Importa uma entrevista, copia para a pasta 'entrevistas' do projeto e salva no banco.
        
        :param filepath: Caminho absoluto do arquivo original a ser importado.
        :param project_root_path: Diretorio raiz do projeto.
        :param sujeito_id: ID do sujeito entrevistado (opcional).
        :return: A instância da Entrevista criada.
        """
        filename = os.path.basename(filepath)
        _, ext = os.path.splitext(filename)
        ext = ext.lower()
        
        if ext == '.txt':
            text = cls.extract_text_from_txt(filepath)
        elif ext == '.docx':
            text = cls.extract_text_from_docx(filepath)
        else:
            raise ValueError(f"Extensão de arquivo não suportada: {ext}. Somente .txt e .docx são suportados.")
        
        # Garante que a pasta de entrevistas exista
        entrevistas_dir = os.path.join(project_root_path, "entrevistas")
        os.makedirs(entrevistas_dir, exist_ok=True)
        
        # Gera nome único na pasta entrevistas para evitar colisão
        dest_filename = filename
        dest_filepath = os.path.join(entrevistas_dir, dest_filename)
        counter = 1
        name_part, ext_part = os.path.splitext(filename)
        while os.path.exists(dest_filepath):
            dest_filename = f"{name_part}_{counter}{ext_part}"
            dest_filepath = os.path.join(entrevistas_dir, dest_filename)
            counter += 1
            
        # Copia o arquivo para a pasta do projeto
        shutil.copy2(filepath, dest_filepath)
        
        # Caminho relativo para armazenar no banco
        caminho_relativo = os.path.join("entrevistas", dest_filename)
        
        # Salva no banco de dados
        db = DatabaseManager.get_instance()
        session = db.get_session()
        
        # Verifica se o sujeito existe se ID fornecido
        if sujeito_id:
            sujeito = session.query(Sujeito).filter(Sujeito.id == sujeito_id).first()
            if not sujeito:
                sujeito_id = None
                
        titulo = name_part.replace("_", " ").title()
        
        nova_entrevista = Entrevista(
            titulo=titulo,
            nome_arquivo_original=filename,
            caminho_relativo_arquivo=caminho_relativo,
            texto_extraido=text,
            data_importacao=datetime.now(),
            sujeito_id=sujeito_id
        )
        
        session.add(nova_entrevista)
        session.commit()
        
        # Recarrega a entrevista com o ID gerado
        session.refresh(nova_entrevista)
        return nova_entrevista
